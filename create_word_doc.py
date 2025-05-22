from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

# Create a new Document
doc = Document()

# Function to add a question and answer to the document
def add_qa(question, answer):
    # Add question
    q_paragraph = doc.add_paragraph()
    q_run = q_paragraph.add_run(question)
    q_run.font.name = 'Calibri'
    q_run.font.size = Pt(13)
    q_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color

    # Add answer
    a_paragraph = doc.add_paragraph()
    a_run = a_paragraph.add_run(answer)
    a_run.font.name = 'Calibri'
    a_run.font.size = Pt(11)
    a_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

# Adding questions and answers
add_qa("Q1.\n", """# 7. Program to check if a number is positive, negative or 0
num = int(input("Enter a Number : "))
if num > 0:
    print("Num is positive")
elif num < 0:
    print("Num is negative")
else:
    print("Number is zero")""")

add_qa("Q2.\n", """# 8. A school decided to replace the desks in three classrooms. Each desk sits two students. Given the number of students in each class, print the smallest possible number of desks that can be purchased.
class1 = int(input("Enter the Number of students in class 1 : "))
class2 = int(input("Enter the Number of students in class 2 : "))
class3 = int(input("Enter the Number of students in class 3 : "))

table1 = (class1 // 2) + (class1 % 2)
table2 = (class2 // 2) + (class2 % 2)
table3 = (class3 // 2) + (class3 % 2)

print(f"Minimum desks in class 1 is {table1}")
print(f"Minimum desks in class 2 is {table2}")
print(f"Minimum desks in class 3 is {table3}")
print("Total Number of Desks required is : ", table1 + table2 + table3)""")

add_qa("Q3.\n", """# 9. H hours, M minutes and S seconds are passed since the midnight (0 ≤ H < 12, 0 ≤ M < 60, 0 ≤ S < 60). Determine the angle (in degrees) of the hour hand on the clock face right now.
H = int(input("Enter Time in Hours (1 to 12): "))
M = int(input("Enter Time in minutes : "))
S = int(input("Enter Time in seconds : "))

# converting hours to degrees
hd = H * 30  # hd means hours degrees

# converting minutes to degrees
md = M * 6  # md means minute degrees

# total degrees
td = hd + md

print(f"{H} Hours, {M} Minutes, {S} Seconds in Clock hand Degrees is : {td} Degrees '{S} Seconds'")""")

add_qa("Q4.\n", """# 10. Given integer coordinates of three vertices of a rectangle whose sides are parallel to the coordinate axes, find the coordinates of the fourth vertex of the rectangle.
import math

xs = []
ys = []

for i in range(0, 3):
    x = int(input(f"Enter Coordinate x{i+1} : "))
    xs.append(x)
    y = int(input(f"Enter Coordinate y{i+1} : "))
    ys.append(y)

if xs[0] == xs[1]:
    x = xs[2]
elif xs[0] == xs[2]:
    x = xs[1]
elif xs[1] == xs[2]:
    x = xs[0]

if ys[0] == ys[1]:
    y = ys[2]
elif ys[0] == ys[2]:
    y = ys[1]
elif ys[1] == ys[2]:
    y = ys[0]

print(f"The Unknown Coordinates are : (X, Y) = ({x}, {y})")""")

add_qa("Q5.\n", """# 11. There was a set of cards with numbers from 1 to N. One of the card is now lost. Determine the number on that lost card given the numbers for the remaining cards.
cards = []

N = int(input("Enter N : "))
sum_list = 0

for i in range(1, N):
    card = int(input("Enter Card Number starting from 1, 2, 3... : "))
    cards.append(card)
    sum_list = cards[i - 1] + sum_list

sum = (N * (N + 1)) / 2

lost_num = sum - sum_list

print(f"The lost Card has Number - '{lost_num}' on it")""")

add_qa("Q6.\n", """# 12. Program to find maximum of three numbers
x = int(input(f"Enter 1st Number : "))
y = int(input(f"Enter 2nd Number : "))
z = int(input(f"Enter 3rd Number : "))

if x >= y and x >= z:
    print("1st Num is the greatest")
elif y >= x and y >= z:
    print("2nd Num is the greatest")
elif z >= x and z >= y:
    print("3rd Num is the greatest")""")

add_qa("Q7.\n", """# 13. Program to check if a year is leap
year = int(input("Enter Year : "))

if year % 100 == 0 and year % 400 == 0:
    print(f"The Year {year} is a Leap Year")
elif year % 4 == 0 and year % 100 != 0:
    print(f"The Year {year} is a Leap Year")
else:
    print(f"Year {year} is NOT Leap Year")""")

add_qa("Q8.\n", """# 14. Program to check if a date is valid
dd, mm, yyyy = map(int, input("Enter Date in 'DD MM YYYY' format : ").split())

m31 = [1, 3, 5, 7, 8, 10, 12]
m30 = [4, 6, 9, 11]
m2 = [2]

if (yyyy % 100 == 0 and yyyy % 400 == 0) or (yyyy % 4 == 0 and yyyy % 100 != 0):
    leap_year = 'true'
else:
    leap_year = 'false'

if dd > 0 and dd <= 31 and (mm in m31):
    d = 'true'
elif dd > 0 and dd <= 30 and (mm in m30):
    d = 'true'
elif dd > 0 and dd <= 29 and leap_year == 'true' and (mm in m2):
    d = 'true'
elif dd > 0 and dd <= 28 and leap_year == 'false' and (mm in m2):
    d = 'true'
else:
    d = 'false'  

if mm > 0 and mm < 13:
    m = 'true'
else:
    m = 'false'
    
if d == 'true' and m == 'true':
    validity = 'valid'
else:
    validity = 'invalid'
    
print(f"The Date Entered {dd}-{mm}-{yyyy} is {validity}")""")

add_qa("Q9.\n", """# 15. Program to find the roots of a quadratic equation
import math

print("The Quadratic Equation is of the Format : ax^2 + bx + c")
a, b, c = map(int, input("Enter the co-efficient of x^2, x, & constant : ").split())

d = ((b * b) - (4 * a * c))
if d >= 0:
    D = math.sqrt(d)
    x1 = (-b + D) / (2 * a)
    x2 = (-b - D) / (2 * a)
    print(f"The Roots of the Equation {a}x^2 + {b}x + {c} are {x1}, {x2}")
else:
    print(f"The Equation {a}x^2 + {b}x + {c} has Imaginary roots")""")

add_qa("Q10.\n", """# 16. Given a string. Delete from it all the characters whose indices are divisible by 3
word = input("Enter any word : ")
new_word = word[0]

for i in range(0, len(word)):
    if i % 3 != 0:
        new_word += word[i]

print(f"After processing the word '{word}' it becomes '{new_word}'")""")

# Adding the next set of questions and answers
add_qa("Q11.\n", """# 17. Given a sequence of integer numbers ending with the number 0. Determine the length of the widest fragment where all the elements are equal to each other.
num_list = [5, 5, 5, 3, 3, 3, 3, 7, 7, 0]
c = 1
frag_len = []

for i in range(len(num_list) - 1):
    if num_list[i] == num_list[i + 1]:
        c += 1
    else:
        frag_len.append(c)
        c = 1

frag_len.sort()
frag_len.reverse()

print(f"The Widest Segment of recurring numbers is of length : {frag_len[0]}")""")

add_qa("Q12.\n", """# 18. In bowling, the player starts with 10 pins at the far end of a lane. The object is to knock all the pins down. For this exercise, the number of pins and balls will vary.
# Given the number of pins N and then the number of balls K to be rolled, followed by K pairs of numbers (one for each ball rolled), determine which pins remain standing after all the balls have been rolled.
# The balls are numbered from 1 to N (inclusive) for this situation.
# The subsequent number pairs, one for each K represent the start to stop (inclusive) positions of the pins that were knocked down with each role.
# Print a sequence of N characters, where "I" represents a pin left standing and "." represents a pin knocked down.
# (Implementation not provided in the original context)""")

add_qa("Q13.\n", """# 19. A timestamp is three numbers: a number of hours, minutes and seconds. Given two timestamps, calculate how many seconds is between them.
hh1, mm1, ss1 = map(int, input("Enter 1st Timestamps in HH:MM:SS format : ").split())
hh2, mm2, ss2 = map(int, input("Enter 2nd Timestamps in HH:MM:SS format : ").split())

seconds = ((hh2 - hh1) * 3600) + ((mm2 - mm1) * 60) + (ss2 - ss1)

print(f"There are {seconds} seconds interval between these two timestamps")""")

add_qa("Q14.\n", """# 20. Program to display first n numbers
n = int(input("Enter any number : "))

for i in range(n):
    print(i + 1, end=" ")""")

add_qa("Q15.\n", """# 21. Program to calculate factorial of a numbers
n = int(input("Enter any number : "))
fact = 1

for i in range(1, n + 1):
    fact = fact * i

print(f"Factorial of {n}! is {fact}")""")

add_qa("Q16.\n", """# 22. Program to display numbers in the reverse order
n = int(input("Enter any number : "))
list_n = []

for i in range(1, n + 1):
    list_n.append(i)

list_n.reverse()

print(f"In reversed : {list_n}")""")

add_qa("Q17.\n", """# 23. Program to check if a number is prime
n = int(input("Enter number to check for prime: "))

for i in range(2, n):
    if n % i != 0:
        c = 1
    else:
        c = 0
        break

if c == 1:
    print(f"The Number {n} is prime")
elif c == 0:
    print(f"{n} is not prime")""")

add_qa("Q18.\n", """# 24. Program to calculate sum and average of first n numbers
n = int(input("Enter any number : "))
sum_n = 0

for i in range(1, n + 1):
    sum_n = sum_n + i

avg = sum_n / n
print(f"For the first {n} Numbers\\nThe Sum is : {sum_n} and Average is : {avg}")""")

add_qa("Q19.\n", """# 25. Program to display first n multiples of a number
n = int(input("Enter any number : "))
multi = int(input("Enter any number : "))

for i in range(multi):
    print((i + 1) * n, end=" ")""")

add_qa("Q20.\n", """# 26. Program to display first n Fibonacci numbers
n = int(input("Enter any number : "))

f = [1, 1]

for i in range(n + 1):
    s = f[i + 1] + f[i]
    f.append(s)

print(f"{n} Fibonacci is : {f}")""")

add_qa("Q21.\n", """# 27. Program to find the sum of digits of a number 
n = int(input("Enter any number : "))

sum_digits = 0
m = n

while m > 0:
    sum_digits += m % 10
    m = int(m / 10)

print(f"Sum of Digits of {n} is {sum_digits}")""")

add_qa("Q22.\n", """# 28. Program to Create and view elements of a list
list_example = []

while True:
    example = input(("Enter anything you'd like to append in the list : "))
    list_example.append(example)
    
    quit_key = input("\\npress 'Q' key to quit : ")
    if quit_key.lower() == 'q':
        break
    
print(f"The contents of list are as follows : {list_example}")""")

add_qa("Q23.\n", """# 29. Program to Create and view elements of a tuple
tuple_example = (
    "curiosity",
    "startup energy",
    "AI enthusiasm",
    "philosophical depth",
    "experimental spirit",
)

print(f"The Example of an TUPLE is : {tuple_example}")""")

add_qa("Q24.\n", """# 30. Program to access List Index and Values
list_example = [
    "curiosity",
    "startup energy",
    "AI enthusiasm",
    "philosophical depth",
    "experimental spirit",
]

i = int(input("Between 1 to 5 ,enter any number which you would like to access in this list : "))

if i > 0 and i < 6:
    print(f"Index '{i}' has '{list_example[i-1]}'")
    print(f"This is the original list : {list_example}")
else:
    print("Please enter between 1 to 5")""")

add_qa("Q25.\n", """# 31. Program to add two Lists
list1 = [
    "curiosity",
    "startup energy",
    "AI enthusiasm",
    "philosophical depth",
    "experimental spirit",
]

list2 = [
    "tech wanderlust",
    "content creation",
    "financial freedom chase",
    "robotics newbie",
    "mentorship vibes",
]

print("List 1 : ", list1)
print("List 2: ", list2)
print(f"\\nThe Addition of lists gives : {list1 + list2}")""")

add_qa("Q26.\n", """# 32. Program to check if a List is Empty or Not
list_example = []

if len(list_example) == 0:
    print("List is empty")""")

add_qa("Q27.\n", """# 33. Program to Find the Largest Number in a List
list_example = []
print("Press '000' to quit : ")

while True:
    example = int(input("Enter any number in any order : "))
    list_example.append(example)

    if example == 000:
        break

list_example.sort()
list_example.reverse()

print(f"The list is : {list_example}")
print(f"\\nThe Largest number of the list is : {list_example[0]}")""")

add_qa("Q28.\n", """# 34. Program to Find the Second Largest Number in a List
list_example = []
print("Press '000' to quit : ")

while True:
    example = int(input("Enter any number in any order : "))
    list_example.append(example)

    if example == 000:
        break

list_example.sort()
list_example.reverse()

print(f"The list is : {list_example}")
print(f"\\nThe Second Largest number of the list is : {list_example[1]}")""")

add_qa("Q29.\n", """# 35. Program to Put Even and Odd elements in a List into Two Different Lists
list_example = []
list_even = []
list_odd = []

print("Press '000' to quit : ")

while True:
    example = int(input("Enter any number in any order : "))

    if example == 000:
        break

    list_example.append(example)

for i in range(len(list_example)):
    if list_example[i] % 2 == 0:
        list_even.append(list_example[i])
    else:
        list_odd.append(list_example[i])

list_even.sort()
list_odd.sort()

print(f"The Unsorted list is {list_example}")
print(f"\\nThe Even sorted list is {list_even}")
print(f"The Odd sorted list is {list_odd}")""")

add_qa("Q30.\n", """# 36. Program to Find all Numbers in a Range which are Perfect Squares and Sum of all Digits in the Number is Less than 10
import math

list_example = []
list_output = []

print("Press '000' to quit : ")

while True:
    example = int(input("Enter any number in any order : "))

    if example == 000:
        break
""")
       
# Save the document
doc.save("Extracted_Code.docx")
